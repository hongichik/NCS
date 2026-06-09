#!/usr/bin/env python3
"""Convert markdown to Word (.docx) with basic formatting."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(strip_md(text))
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text.strip()


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        else:
            p.add_run(strip_md(part))
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(strip_md(line[2:]), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md(line[3:]), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(strip_md(line[4:]), level=2)
            i += 1
            continue

        if line.strip().startswith("> "):
            add_rich_paragraph(doc, line.strip()[2:], style="Intense Quote")
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            headers = parse_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            ncols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = "Table Grid"
            for j, h in enumerate(headers):
                set_cell_text(table.rows[0].cells[j], h, bold=True)
            for r_idx, row in enumerate(rows):
                for c_idx in range(ncols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)
            doc.add_paragraph()
            continue

        if line.strip().startswith("- "):
            add_rich_paragraph(doc, line.strip()[2:], style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s", "", line.strip()), style="List Number")
            i += 1
            continue

        if line.strip().startswith("- [ ]"):
            add_rich_paragraph(doc, "☐ " + line.strip()[5:].strip())
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_rich_paragraph(doc, line)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("TONG_HOP_RetailRocket_CatSA.md")
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    convert(md, out)
